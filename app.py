import os
import sys
import types

# --- BULLETPROOF RUNTIME MEMORY MOCK PATCH ---
# Force-create a dummy pkg_resources module in system memory to satisfy legacy CrewAI telemetry
IMAGINARY_PKG = types.ModuleType("pkg_resources")
IMAGINARY_PKG.declare_namespace = lambda name: None
IMAGINARY_PKG.get_distribution = lambda name: types.SimpleNamespace(version="0.28.8")
sys.modules["pkg_resources"] = IMAGINARY_PKG
# ──────────────────────────────────────

import streamlit as st
import io
from docx import Document
import threading
from streamlit.runtime.scriptrunner import add_script_run_ctx
from docx import Document

# Add source directory path and import your crew module
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))
from src.job_crew.crew import JobApplicationCrew


st.set_page_config(page_title="AI Application Tailoring Crew", page_icon="💼", layout="wide")

st.title("💼 Multi-Agent Job Tailoring System")
st.caption("Coordinate autonomous agents to evaluate listings, profiles, and deliver targeted interview briefs.")

# Initialize background worker states
if "strategy_data" not in st.session_state:
    st.session_state.strategy_data = None
if "resume_data" not in st.session_state:
    st.session_state.resume_data = None
if "questions_data" not in st.session_state:
    st.session_state.questions_data = None
if "talking_points_data" not in st.session_state:
    st.session_state.talking_points_data = None
if "crew_running" not in st.session_state:
    st.session_state.crew_running = False
if "crew_result" not in st.session_state:
    st.session_state.crew_result = None

import os
import io
import streamlit as st
from docx import Document  # python-docx is already installed in your environment

def convert_md_to_docx(md_text):
    """Converts raw markdown text into a downloadable Word Document (.docx) bytes stream"""
    doc = Document()
    # Simple line-by-line parsing for basic markdown compatibility (headers and paragraphs)
    for line in md_text.split('\n'):
        stripped = line.strip()
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped:
            doc.add_paragraph(stripped)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def render_download_buttons(file_content, base_filename):
    """Renders side-by-side download buttons for TXT and Word documents"""
    col1, col2 = st.columns(2)
    
    with col1:
        st.download_button(
            label="📥 Download as TXT",
            data=file_content,
            file_name=f"{base_filename}.txt",
            mime="text/plain",
            key=f"txt_{base_filename}"
        )
        
    with col2:
        docx_bytes = convert_md_to_docx(file_content)
        st.download_button(
            label="📥 Download as Word (DOCX)",
            data=docx_bytes,
            file_name=f"{base_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"docx_{base_filename}"
        )
        
def run_crew_async(inputs):
    try:
        # Run the crew without touching st.session_state inside here
        result = JobApplicationCrew().crew().kickoff(inputs=inputs)
        return result
    except Exception as e:
        print(f"Error running crew: {e}")

# Layout Forms
with st.sidebar:
    st.header("Pipeline Inputs")
    job_url = st.text_input("Job Posting URL", value="https://jobs.lever.co/AIFund/6c82e23e-d954-4dd8-a734-c0c2c5ee00f1")
    
    st.markdown("---")
    st.subheader("Candidate Information")
    
    # 1. Flexible Input Strategy Selector
    profile_source = st.radio(
        "Choose profile background source:",
        ["Provide Profile URL (LinkedIn / GitHub / Portfolio)", "Upload Profile Document (TXT / MD)"]
    )
    
    #
    # 1. Initialize session state keys at the very top of your input section
# =====================================================================
# STEP 1: INITIALIZE ALL STATE KEYS AND VARIABLES AT THE VERY TOP
# =====================================================================
if "uploaded_text_content" not in st.session_state:
    st.session_state.uploaded_text_content = ""

if "strategy_data" not in st.session_state:
    st.session_state.strategy_data = None
if "resume_data" not in st.session_state:
    st.session_state.resume_data = None
if "questions_data" not in st.session_state:
    st.session_state.questions_data = None
if "talking_points_data" not in st.session_state:
    st.session_state.talking_points_data = None

# Pre-define fallback variables to guarantee they ALWAYS exist
profile_url_input = ""
uploaded_text_content = ""
text_summary = ""


# =====================================================================
# STEP 2: RENDER THE USER INTERFACE INPUT FIELDS
# =====================================================================
# Main Bio/Resume Text Input Area
text_summary = st.text_area(
    "Paste Personal Bio or Background Summary", 
    value="Experienced Software Engineer with a background in full-stack systems..."
)

# Profile Source Selection & File Uploader Block
if profile_source == "Provide Profile URL (LinkedIn / GitHub / Portfolio)":
    profile_url_input = st.text_input("Profile URL", value="https://github.com/joaomdmoura")
else:
    uploaded_file = st.file_uploader(
        "Upload background profile or existing resume", 
        type=["txt", "md"],
        key="profile_file_uploader"
    )
    
    if uploaded_file is not None:
        try:
            file_bytes = uploaded_file.read()
            text_content = file_bytes.decode("utf-8")
            
            # Lock into memory safely
            if st.session_state.uploaded_text_content != text_content:
                st.session_state.uploaded_text_content = text_content
                st.success("File uploaded and saved to memory successfully!")
        except Exception as e:
            st.error(f"Error reading file stream: {str(e)}")

    # Pull latest data from stable session memory
    uploaded_text_content = st.session_state.uploaded_text_content

st.markdown("---")


# =====================================================================
# STEP 3: THE ACTION TRIGGER BUTTON
# =====================================================================
submit_btn = st.button("Kickoff Crew Agents", disabled=st.session_state.crew_running)


# =====================================================================
# STEP 4: THE CONTROLLER ACTION BLOCK (READS COMPLETED STEP 1 & 2 DATA)
# =====================================================================
if submit_btn and not st.session_state.crew_running:
    st.session_state.crew_running = True
    st.session_state.crew_result = None
    
    # Reset stored outputs on a new fresh run
    st.session_state.strategy_data = None
    st.session_state.resume_data = None
    st.session_state.questions_data = None
    st.session_state.talking_points_data = None
    
    # Consolidate text content based on what the user provided
    final_writeup = text_summary
    if uploaded_text_content:
        final_writeup += f"\n\n--- EMBEDDED ATTACHMENT PROFILE CONTENT ---\n{uploaded_text_content}"
    
    # ✅ NOW crew_inputs CANNOT THROW A NAMEERROR BECAUSE EVERYTHING ABOVE IS LOADED
    crew_inputs = {
        'job_posting_url': job_url, # Ensure job_url widget is placed above this section
        'profile_urls': profile_url_input if profile_url_input else "No URL provided; check text/attachment profile data.",
        'personal_writeup': final_writeup
    }

    # ─── RUN THE CREW NATIVELY HERE ───
    with st.spinner("Formulating strategy... The crew is analyzing your data live."):
        try:
            # Clean up residual layout files
            for f in ['strategy_output.md', 'tailored_resume.md', 'questions_output.md', 'talking_points_output.md']:
                if os.path.exists(f):
                    os.remove(f)

            # Run CrewAI Framework
            crew_instance = JobApplicationCrew().crew()
            result = crew_instance.kickoff(inputs=crew_inputs)
            st.session_state.crew_result = result
            
            # Load into persistent memory instantly
            if os.path.exists('strategy_output.md'):
                with open('strategy_output.md', 'r', encoding='utf-8') as f:
                    st.session_state.strategy_data = f.read()
            if os.path.exists('tailored_resume.md'):
                with open('tailored_resume.md', 'r', encoding='utf-8') as f:
                    st.session_state.resume_data = f.read()
            if os.path.exists('questions_output.md'):
                with open('questions_output.md', 'r', encoding='utf-8') as f:
                    st.session_state.questions_data = f.read()
            if os.path.exists('talking_points_output.md'):
                with open('talking_points_output.md', 'r', encoding='utf-8') as f:
                    st.session_state.talking_points_data = f.read()

            st.success("Analysis Complete!")
            
        except Exception as e:
            st.error(f"An execution error occurred: {str(e)}")
        finally:
            st.session_state.crew_running = False

# ─── DISPLAY THE TABS RENDERING OUTSIDE THE BUTTON RUN BLOCK ───
# This ensures that even when downloads rerun the script, the content persists!
if st.session_state.strategy_data or st.session_state.resume_data:
    tab1, tab2, tab3, tab4 = st.tabs([
        "🎯 Tailored Strategy & Prep", 
        "📄 Custom Resume",
        "❓ Expected Interview Questions", 
        "📋 Custom Talking Points"
    ])
    
    with tab1:
        st.subheader("Application Strategy & Profile Alignment")
        if st.session_state.strategy_data:
            st.markdown(st.session_state.strategy_data)
            st.divider()
            render_download_buttons(st.session_state.strategy_data, "application_strategy")
        else:
            st.info("No strategy data available.")
            
    with tab2:
        st.subheader("Tailored Resume Content")
        if st.session_state.resume_data:
            st.markdown(st.session_state.resume_data)
            st.divider()
            render_download_buttons(st.session_state.resume_data, "tailored_resume")
        else:
            st.info("The resume customization task did not compile a separate file.")
            
    with tab3:
        st.subheader("Targeted Interview Preparation Questions")
        if st.session_state.questions_data:
            st.markdown(st.session_state.questions_data)
            st.divider()
            render_download_buttons(st.session_state.questions_data, "interview_questions")
        else:
            st.info("Interview questions preparation was incomplete.")
            
    with tab4:
        st.subheader("Key Interview Talking Points")
        if st.session_state.talking_points_data:
            st.markdown(st.session_state.talking_points_data)
            st.divider()
            render_download_buttons(st.session_state.talking_points_data, "interview_talking_points")
        else:
            st.info("Strategic talking points document was incomplete.")
      